from docx import Document

document = Document("iob.docx")
tables = document.tables
for col in tables[0].columns:
    print(col.cells[0].paragraphs[0].text,end="  ")
print(" \n ")
for col in tables[0].columns:
    print(col.cells[1].paragraphs[0].text,end="  ")
print(" \n ")
for col in tables[0].columns:
    print(col.cells[2].paragraphs[0].text,end="  ")
print(" \n ")
for col in tables[0].columns:
    print(col.cells[3].paragraphs[0].text,end="  ")
